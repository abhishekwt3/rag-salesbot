# main.py - SaaS Multi-tenant RAG Chatbot API (Enhanced Scraper)
import os
import logging
from typing import List
from datetime import datetime, timezone
from pathlib import Path
import tempfile
import docx
import time as time_module

# Load .env file
from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, HTTPException, BackgroundTasks, Depends, UploadFile, File, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, RedirectResponse
from typing import List, Dict
import uvicorn
import uuid

# Import database models and utilities
#from models import User, KnowledgeBase, ChatWidget, WidgetConversation, get_db, init_database
from auth import get_current_user, hash_password, create_access_token, authenticate_user
from schemas import (
    UserCreate, UserLogin, Token, 
    KnowledgeBaseCreate,
    WebsiteConfig, ProcessingResponse,
    QueryRequest, QueryResponse,
    MessageResponse, ChatWidgetCreate, ChatWidgetResponse, ChatWidgetUpdate,
    WidgetChatRequest, WidgetChatResponse, WidgetConfigResponse, DocumentUploadResponse
)
from sqlalchemy.orm import Session
from models import (
    User, KnowledgeBase, ChatWidget, WidgetConversation, get_db, init_database,
    Subscription, Payment, SubscriptionPlan, SubscriptionStatus, PaymentStatus
)
from subscription_schemas import *
from subscription_middleware import (
    subscription_manager, get_user_subscription, check_subscription_active,
    check_can_create_kb, validate_chunk_addition,
    validate_kb_creation, KBUsageTracker
)
from razorpay_utils import razorpay_manager

# Import the enhanced vector store and scraper
from saas_embeddings import MemcacheS3VectorStore
from simple_scraper import EnhancedSimpleScraper as WebScraper

# Import Google OAuth
from google_oauth import GoogleOAuth, get_google_oauth_endpoints

import httpx

def validate_environment():
    """Validate required environment variables"""
    required_vars = [
        'QDRANT_URL',
        'GOOGLE_API_KEY'
    ]
    
    optional_vars = [
        'QDRANT_API_KEY'  # Optional for local Qdrant
    ]
    
    missing_vars = []
    for var in required_vars:
        if not os.getenv(var):
            missing_vars.append(var)
    
    if missing_vars:
        logger.warning(f"Missing required environment variables: {missing_vars}")
        logger.warning("Some features may not work properly")
    
    # Log optional variables
    for var in optional_vars:
        if not os.getenv(var):
            logger.info(f"Optional environment variable not set: {var}")
    
    return len(missing_vars) == 0

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Validate environment and initialize database
validate_environment()
init_database()

app = FastAPI(title="SaaS RAG Chatbot API", version="4.1.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Enhanced Vector store manager with scraper integration
class SaaSVectorManager:
    def __init__(self):
        self.vector_stores = {}  # user_id -> knowledge_base_id -> vector_store
        # Initialize enhanced scraper with optimized settings
        self.scraper = WebScraper(
            max_retries=3,
            host_max_concurrent=2,  # Conservative for shared hosting
            host_min_interval_ms=500,  # Slower for politeness
            respect_robots=True,
            enable_resource_blocking=True,  # 3-5x performance boost
            on_result=self._on_scrape_result
        )
        
    def _on_scrape_result(self, tenant_id: str, page):
        """Callback when a page is successfully scraped"""
        logger.info(f"📄 Scraped: {page.final_url} ({page.framework}, {page.to_dict()['word_count']} words) for {tenant_id}")
        
    def get_vector_store(self, user_id: str, knowledge_base_id: str) -> MemcacheS3VectorStore:
        if user_id not in self.vector_stores:
            self.vector_stores[user_id] = {}
        
        if knowledge_base_id not in self.vector_stores[user_id]:
            self.vector_stores[user_id][knowledge_base_id] = MemcacheS3VectorStore(
                user_id=user_id,
                knowledge_base_id=knowledge_base_id
            )
        
        return self.vector_stores[user_id][knowledge_base_id]
    
    def clear_vector_store(self, user_id: str, knowledge_base_id: str):
        """Clear a specific vector store"""
        if user_id in self.vector_stores:
            if knowledge_base_id in self.vector_stores[user_id]:
                # Clear the vector store data
                self.vector_stores[user_id][knowledge_base_id].clear_data()
                # Remove from memory
                del self.vector_stores[user_id][knowledge_base_id]
    
    async def scrape_websites(self, jobs: List[Dict]) -> Dict[str, List[Dict]]:
        """
        Enhanced multi-tenant website scraping
        jobs: [{"tenant_id": "user_id:kb_id", "urls": [...]}]
        returns: {tenant_id: [page_dict, ...]}
        """
        return await self.scraper.scrape_multi_tenant(jobs)

# Initialize vector manager
vector_manager = SaaSVectorManager()

# Google OAuth setup
google_oauth = GoogleOAuth()
oauth_endpoints = get_google_oauth_endpoints()

# Health check
@app.get("/")
async def root():
    return {"message": "Salesdok Assistant API is running", "version": "4.1.0", "scraper": "enhanced", "vector_db": "weaviate"}

@app.get("/health")
async def health():
    return {"status": "healthy", "scraper_features": ["multi-tenant", "spa-support", "resource-blocking", "robots-txt"]}

# Auth endpoints
@app.post("/auth/register", response_model=Token)
async def register(user_data: UserCreate, db: Session = Depends(get_db)):
    # Check if user already exists
    if db.query(User).filter(User.email == user_data.email).first():
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create new user
    hashed_password = hash_password(user_data.password)
    db_user = User(
        email=user_data.email,
        hashed_password=hashed_password,
        full_name=user_data.full_name
    )
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    
    # Create access token
    access_token = create_access_token(data={"sub": str(db_user.id)})
    
    logger.info(f"New user registered: {user_data.email}")
    
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/auth/login", response_model=Token)
async def login(user_data: UserLogin, db: Session = Depends(get_db)):
    user = authenticate_user(user_data.email, user_data.password, db)
    
    if not user:
        raise HTTPException(status_code=401, detail="Incorrect email or password")
    
    if not user.is_active:
        raise HTTPException(status_code=401, detail="User account is disabled")
    
    access_token = create_access_token(data={"sub": str(user.id)})
    
    logger.info(f"User logged in: {user_data.email}")
    
    return {"access_token": access_token, "token_type": "bearer"}

# Google OAuth endpoints (unchanged)
@app.get("/auth/google")
async def google_auth():
    """Initiate Google OAuth flow"""
    if not os.getenv("GOOGLE_CLIENT_ID"):
        raise HTTPException(status_code=501, detail="Google OAuth not configured. Please set GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET environment variables.")
    
    auth_url = google_oauth.get_auth_url()
    return RedirectResponse(url=auth_url)

@app.get("/auth/google/callback")
async def google_callback(
    code: str = Query(...),
    state: str = Query(None),
    error: str = Query(None),
    db: Session = Depends(get_db)
):
    """Handle Google OAuth callback"""
    
    if error:
        logger.error(f"Google OAuth error: {error}")
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        return RedirectResponse(url=f"{frontend_url}?error=oauth_error")
    
    if not code:
        raise HTTPException(status_code=400, detail="Authorization code not provided")
    
    try:
        # Exchange code for token
        token_data = await google_oauth.exchange_code_for_token(code)
        access_token = token_data.get("access_token")
        
        if not access_token:
            raise HTTPException(status_code=400, detail="No access token received")
        
        # Get user info from Google
        user_info = await google_oauth.get_user_info(access_token)
        
        email = user_info.get("email")
        name = user_info.get("name", "")
        
        if not email:
            raise HTTPException(status_code=400, detail="Email not provided by Google")
        
        # Check if user exists
        existing_user = db.query(User).filter(User.email == email).first()
        
        if existing_user:
            user = existing_user
            logger.info(f"Existing user logged in via Google: {email}")
        else:
            # Create new user
            user = User(
                email=email,
                full_name=name,
                hashed_password="oauth_user",
            )
            
            db.add(user)
            db.commit()
            db.refresh(user)
            
            logger.info(f"New user created via Google OAuth: {email}")
        
        # Create JWT token
        jwt_token = create_access_token(data={"sub": str(user.id)})
        
        # Redirect to frontend with token
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        return RedirectResponse(
            url=f"{frontend_url}?token={jwt_token}&email={email}&name={name}"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Google OAuth callback error: {e}")
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        return RedirectResponse(url=f"{frontend_url}?error=oauth_failed")

@app.get("/auth/me")
async def get_current_user_info(current_user: User = Depends(get_current_user)):
    return {
        "id": str(current_user.id),
        "email": current_user.email,
        "full_name": current_user.full_name,
        "created_at": current_user.created_at
    }

@app.get("/subscription/plans", response_model=PlansListResponse)
async def get_available_plans():
    """Get all available subscription plans"""
    plans = []
    
    for plan_type in SubscriptionPlan:
        plan_details = Subscription.get_plan_details(plan_type)
        
        # Define features for each plan
        features_map = {
            SubscriptionPlan.BASIC: [
                "5 knowledge bases",
                "200 data chunks total",
                "Email support",
                "Standard branding",
                "Basic analytics"
            ],
            SubscriptionPlan.PRO: [
                "30 knowledge bases", 
                "1,500 data chunks total",
                "Priority support",
                "Custom branding",
                "Advanced analytics",
                "API access",
                "Integrations"
            ],
            SubscriptionPlan.ENTERPRISE: [
                "Unlimited knowledge bases",
                "Unlimited data chunks", 
                "24/7 dedicated support",
                "White-label solution",
                "Enterprise analytics",
                "SSO & SAML",
                "Custom integrations",
                "SLA guarantee"
            ]
        }
        
        plan_info = PlanInfoResponse(
            plan=plan_type.value,
            name=plan_details["name"],
            amount_usd=plan_details["amount_usd"],
            amount_inr=plan_details["amount_inr"],
            max_knowledge_bases=plan_details["max_knowledge_bases"],
            max_total_chunks=plan_details["max_total_chunks"],
            description=plan_details["description"],
            features=features_map[plan_type]
        )
        plans.append(plan_info)
    
    return PlansListResponse(plans=plans)

@app.get("/subscription/current", response_model=SubscriptionResponse)
async def get_current_subscription(
    subscription: Subscription = Depends(get_user_subscription)
):
    """Get user's current subscription details"""
    remaining_chunks = subscription.get_remaining_chunks()
    
    return SubscriptionResponse(
        id=subscription.id,
        plan=subscription.plan.value,
        status=subscription.status.value,
        amount=subscription.amount,
        currency=subscription.currency,
        billing_cycle=subscription.billing_cycle,
        max_knowledge_bases=subscription.max_knowledge_bases,
        max_total_chunks=subscription.max_total_chunks,
        current_chunk_usage=subscription.current_chunk_usage,
        current_kb_count=subscription.current_kb_count,
        remaining_chunks=remaining_chunks,
        trial_end=subscription.trial_end,
        current_period_start=subscription.current_period_start,
        current_period_end=subscription.current_period_end,
        created_at=subscription.created_at,
        is_trial_active=subscription.is_trial_active()
    )

@app.get("/subscription/usage", response_model=UsageResponse)
async def get_usage_details(
    subscription: Subscription = Depends(get_user_subscription)
):
    """Get detailed usage information"""
    remaining_chunks = subscription.get_remaining_chunks()
    
    return UsageResponse(
        current_chunk_usage=subscription.current_chunk_usage,
        max_total_chunks=subscription.max_total_chunks,
        remaining_chunks=remaining_chunks,
        current_kb_count=subscription.current_kb_count,
        max_knowledge_bases=subscription.max_knowledge_bases,
        can_create_kb=subscription.can_create_knowledge_base(),
        plan=subscription.plan.value,
        status=subscription.status.value
    )

@app.post("/subscription/create", response_model=SubscriptionCreateResponse)
async def create_subscription(
    subscription_data: SubscriptionCreateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new subscription"""
    if not razorpay_manager.is_available():
        raise HTTPException(
            status_code=501,
            detail="Payment processing is not available. Please contact support."
        )
    
    # Check if user already has an active subscription
    existing_subscription = db.query(Subscription).filter(
        Subscription.user_id == current_user.id,
        Subscription.status.in_([SubscriptionStatus.ACTIVE, SubscriptionStatus.TRIALING])
    ).first()
    
    if existing_subscription:
        raise HTTPException(
            status_code=400,
            detail="You already have an active subscription. Use the upgrade endpoint to change plans."
        )
    
    try:
        # Create Razorpay customer
        customer = razorpay_manager.create_customer(
            email=current_user.email,
            name=current_user.full_name
        )
        
        if not customer:
            raise HTTPException(status_code=500, detail="Failed to create customer")
        
        # Create Razorpay plan
        plan = razorpay_manager.create_plan(
            subscription_data.plan,
            subscription_data.currency
        )
        
        if not plan:
            raise HTTPException(status_code=500, detail="Failed to create plan")
        
        # Get plan details
        plan_details = Subscription.get_plan_details(subscription_data.plan)
        amount_key = f"amount_{subscription_data.currency.lower()}"
        amount = plan_details[amount_key]
        
        # Create payment link for immediate payment
        payment_link = razorpay_manager.create_payment_link(
            amount=amount,
            currency=subscription_data.currency,
            customer_email=current_user.email,
            description=f"{plan_details['name']} - Monthly Subscription"
        )
        
        # Create subscription in database
        subscription = Subscription(
            user_id=current_user.id,
            plan=subscription_data.plan,
            status=SubscriptionStatus.PENDING,
            razorpay_customer_id=customer['id'],
            razorpay_plan_id=plan['id'],
            amount=amount,
            currency=subscription_data.currency,
            billing_cycle=subscription_data.billing_cycle,
            max_knowledge_bases=plan_details["max_knowledge_bases"],
            max_total_chunks=plan_details["max_total_chunks"],
            current_chunk_usage=0,
            current_kb_count=0
        )
        
        db.add(subscription)
        db.commit()
        db.refresh(subscription)
        
        logger.info(f"✅ Created subscription for user {current_user.email}: {subscription_data.plan.value}")
        
        return SubscriptionCreateResponse(
            subscription_id=subscription.id,
            razorpay_subscription_id=subscription.razorpay_subscription_id,
            payment_link=payment_link.get('short_url') if payment_link else None,
            trial_end=subscription.trial_end,
            status=subscription.status.value,
            message="Subscription created successfully. Please complete payment to activate."
        )
        
    except Exception as e:
        logger.error(f"❌ Error creating subscription: {e}")
        raise HTTPException(status_code=500, detail="Failed to create subscription")

@app.post("/subscription/payment/verify", response_model=PaymentResponse)
async def verify_payment(
    payment_data: PaymentRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Verify payment and activate subscription"""
    
    # Get subscription
    subscription = db.query(Subscription).filter(
        Subscription.id == payment_data.subscription_id,
        Subscription.user_id == current_user.id
    ).first()
    
    if not subscription:
        raise HTTPException(status_code=404, detail="Subscription not found")
    
    # Verify payment signature
    is_valid = razorpay_manager.verify_payment_signature(
        payment_data.razorpay_payment_id,
        payment_data.razorpay_order_id,
        payment_data.razorpay_signature
    )
    
    if not is_valid:
        raise HTTPException(status_code=400, detail="Invalid payment signature")
    
    try:
        # Get payment details from Razorpay
        payment_details = razorpay_manager.get_payment_details(payment_data.razorpay_payment_id)
        
        if not payment_details or payment_details.get('status') != 'captured':
            raise HTTPException(status_code=400, detail="Payment not successful")
        
        # Create payment record
        payment = Payment(
            subscription_id=subscription.id,
            user_id=current_user.id,
            razorpay_payment_id=payment_data.razorpay_payment_id,
            razorpay_order_id=payment_data.razorpay_order_id,
            razorpay_signature=payment_data.razorpay_signature,
            amount=payment_details['amount'],
            currency=payment_details['currency'],
            status=PaymentStatus.COMPLETED,
            payment_method=payment_details.get('method'),
            paid_at=datetime.now(timezone.utc),
            billing_period_start=datetime.now(timezone.utc),
            billing_period_end=datetime.now(timezone.utc).replace(month=datetime.now(timezone.utc).month + 1)
        )
        
        db.add(payment)
        
        # Activate subscription
        subscription.status = SubscriptionStatus.ACTIVE
        subscription.current_period_start = datetime.now(timezone.utc)
        subscription.current_period_end = datetime.now(timezone.utc).replace(month=datetime.now(timezone.utc).month + 1)
        
        db.commit()
        
        logger.info(f"✅ Payment verified and subscription activated for user {current_user.email}")
        
        return PaymentResponse(
            status="success",
            message="Payment verified and subscription activated",
            payment_id=payment.id,
            subscription_status=subscription.status.value
        )
        
    except Exception as e:
        logger.error(f"❌ Error verifying payment: {e}")
        raise HTTPException(status_code=500, detail="Failed to verify payment")

@app.post("/webhooks/razorpay")
async def razorpay_webhook(
    webhook_data: WebhookRequest,
    db: Session = Depends(get_db)
):
    """Handle Razorpay webhooks"""
    
    logger.info(f"📡 Received Razorpay webhook: {webhook_data.event}")
    
    try:
        result = razorpay_manager.handle_webhook(webhook_data.dict())
        
        # Handle specific webhook events
        if result["event"] == "payment.captured":
            # Update payment status
            payment_id = result.get("payment_id")
            if payment_id:
                payment = db.query(Payment).filter(
                    Payment.razorpay_payment_id == payment_id
                ).first()
                
                if payment:
                    payment.status = PaymentStatus.COMPLETED
                    payment.paid_at = datetime.now(timezone.utc)
                    
                    # Activate subscription
                    subscription = payment.subscription
                    subscription.status = SubscriptionStatus.ACTIVE
                    
                    db.commit()
        
        elif result["event"] == "subscription.cancelled":
            # Cancel subscription
            subscription_id = result.get("subscription_id")
            if subscription_id:
                subscription = db.query(Subscription).filter(
                    Subscription.razorpay_subscription_id == subscription_id
                ).first()
                
                if subscription:
                    subscription.status = SubscriptionStatus.CANCELED
                    subscription.canceled_at = datetime.now(timezone.utc)
                    db.commit()
        
        return WebhookResponse(status="processed", message="Webhook processed successfully")
        
    except Exception as e:
        logger.error(f"❌ Error processing webhook: {e}")
        return WebhookResponse(status="error", message=str(e))

# Knowledge base endpoints
@app.post("/knowledge-bases")
async def create_knowledge_base_with_limits(
    kb_data: KnowledgeBaseCreate,
    current_user: User = Depends(get_current_user),
    subscription: Subscription = Depends(check_can_create_kb),
    db: Session = Depends(get_db)
):
    """Create knowledge base with subscription limits"""
    
    with KBUsageTracker(current_user.id, db) as tracker:
        knowledge_base = KnowledgeBase(
            user_id=current_user.id,
            name=kb_data.name,
            description=kb_data.description
        )
        db.add(knowledge_base)
        db.commit()
        db.refresh(knowledge_base)
        
        tracker.mark_created()
        
        logger.info(f"Knowledge base created with limits: {kb_data.name} for user {current_user.email}")
        
        return {
            "id": str(knowledge_base.id),
            "name": knowledge_base.name,
            "description": knowledge_base.description,
            "status": knowledge_base.status,
            "created_at": knowledge_base.created_at,
            "subscription_info": {
                "current_kb_count": subscription.current_kb_count + 1,
                "max_knowledge_bases": subscription.max_knowledge_bases,
                "plan": subscription.plan.value
            }
        }

@app.get("/knowledge-bases")
async def list_knowledge_bases(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    knowledge_bases = db.query(KnowledgeBase).filter(
        KnowledgeBase.user_id == current_user.id
    ).all()
    
    return [
        {
            "id": str(kb.id),
            "name": kb.name,
            "description": kb.description,
            "status": kb.status,
            "total_chunks": kb.total_chunks,
            "last_updated": kb.last_updated,
            "created_at": kb.created_at
        }
        for kb in knowledge_bases
    ]

@app.delete("/knowledge-bases/{knowledge_base_id}")
async def delete_knowledge_base(
    knowledge_base_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Verify ownership
    knowledge_base = db.query(KnowledgeBase).filter(
        KnowledgeBase.id == knowledge_base_id,
        KnowledgeBase.user_id == current_user.id
    ).first()
    
    if not knowledge_base:
        raise HTTPException(status_code=404, detail="Knowledge base not found")
    
    # Clear vector store in Weaviate
    vector_manager.clear_vector_store(str(current_user.id), knowledge_base_id)
    
    db.delete(knowledge_base)
    db.commit()
    
    logger.info(f"Knowledge base deleted: {knowledge_base.name}")
    
    return MessageResponse(message="Knowledge base deleted successfully")

# Update website processing endpoint to check chunk limits
@app.post("/knowledge-bases/{knowledge_base_id}/process-website")
async def process_website_with_limits(
    knowledge_base_id: str,
    config: WebsiteConfig,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    subscription: Subscription = Depends(check_subscription_active),
    db: Session = Depends(get_db)
):
    """Process website with chunk limit validation"""
    
    # Verify ownership
    knowledge_base = db.query(KnowledgeBase).filter(
        KnowledgeBase.id == knowledge_base_id,
        KnowledgeBase.user_id == current_user.id
    ).first()
    
    if not knowledge_base:
        raise HTTPException(status_code=404, detail="Knowledge base not found")
    
    # Estimate chunk usage (rough estimate: 1 page ≈ 5-10 chunks)
    estimated_chunks = config.max_pages * 8  # Conservative estimate
    
    # Check if user can add estimated chunks
    validation = validate_chunk_addition(current_user.id, estimated_chunks, db)
    if not validation["can_add"]:
        raise HTTPException(
            status_code=403,
            detail=f"Estimated {estimated_chunks} chunks would exceed your limit. {validation['message']} Consider upgrading your plan."
        )
    
    # Update status to processing
    knowledge_base.status = "processing"
    knowledge_base.website_url = str(config.url)
    db.commit()
    
    # Add background task
    background_tasks.add_task(
        process_website_background_with_limits,
        str(current_user.id),
        knowledge_base_id,
        config
    )
    
    logger.info(f"🚀 Website processing started with limits for KB {knowledge_base_id}: {config.url}")
    
    return ProcessingResponse(
        message="Website processing started with subscription limits",
        knowledge_base_id=knowledge_base_id,
        status="processing"
    )


async def process_website_background_with_limits(user_id: str, knowledge_base_id: str, config: WebsiteConfig):
    """Enhanced background task with subscription limit enforcement - FIXED"""
    from models import SessionLocal
    db = SessionLocal()
    
    try:
        logger.info(f"🚀 Starting website processing for KB {knowledge_base_id}, URL: {config.url}")
        
        # Get user's subscription
        subscription = db.query(Subscription).filter(
            Subscription.user_id == user_id
        ).first()
        
        if not subscription:
            logger.error(f"❌ No subscription found for user {user_id}")
            # Update KB status to error
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
            if kb:
                kb.status = "error"
                db.commit()
            return
        
        # Get vector store for this user and knowledge base
        vector_store = vector_manager.get_vector_store(user_id, knowledge_base_id)
        
        # Prepare URLs for scraping
        urls = []
        main_url = str(config.url)
        urls.append(main_url)
        
        # If additional URLs are provided in config, add them
        if hasattr(config, 'additional_urls') and config.additional_urls:
            urls.extend([str(url) for url in config.additional_urls])
        
        # Create multi-tenant job for enhanced scraper
        tenant_id = f"{user_id}:{knowledge_base_id}"
        scraping_jobs = [{
            "tenant_id": tenant_id,
            "urls": urls
        }]
        
        logger.info(f"📄 Starting enhanced scraping for {len(urls)} URLs...")
        logger.info(f"📄 URLs to scrape: {urls}")
        
        # Scrape using enhanced multi-tenant scraper
        try:
            results = await vector_manager.scrape_websites(scraping_jobs)
            pages = results.get(tenant_id, [])
            logger.info(f"📄 Scraper returned {len(pages)} pages")
        except Exception as scrape_error:
            logger.error(f"❌ Scraping failed: {scrape_error}")
            # Update KB status to error
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
            if kb:
                kb.status = "error"
                db.commit()
            return
        
        if not pages:
            logger.warning(f"⚠️ No pages scraped for KB {knowledge_base_id}")
            logger.warning(f"⚠️ Scraping results: {results}")
            
            # Try fallback simple scraping approach
            logger.info(f"🔄 Trying fallback scraping approach...")
            try:
                import requests
                from bs4 import BeautifulSoup
                
                # Simple fallback scraping
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                
                fallback_pages = []
                for url in urls[:config.max_pages]:  # Respect max_pages limit
                    try:
                        response = requests.get(url, headers=headers, timeout=10)
                        if response.status_code == 200:
                            soup = BeautifulSoup(response.content, 'html.parser')
                            
                            # Extract title
                            title = ""
                            if soup.title:
                                title = soup.title.string.strip() if soup.title.string else ""
                            
                            # Extract text content
                            # Remove script and style elements
                            for script in soup(["script", "style"]):
                                script.decompose()
                            
                            # Get text
                            text = soup.get_text()
                            
                            # Clean up text
                            lines = (line.strip() for line in text.splitlines())
                            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                            text = ' '.join(chunk for chunk in chunks if chunk)
                            
                            if len(text.strip()) > 100:  # Only add if meaningful content
                                page_data = {
                                    "final_url": url,
                                    "title": title,
                                    "text": text,
                                    "word_count": len(text.split()),
                                    "status": response.status_code,
                                    "framework": "fallback",
                                    "hash": str(hash(text))
                                }
                                fallback_pages.append(page_data)
                                logger.info(f"✅ Fallback scraped: {url} ({len(text)} chars)")
                    
                    except Exception as e:
                        logger.error(f"❌ Fallback scraping failed for {url}: {e}")
                        continue
                
                pages = fallback_pages
                logger.info(f"📄 Fallback scraping got {len(pages)} pages")
                
            except Exception as fallback_error:
                logger.error(f"❌ Fallback scraping failed: {fallback_error}")
        
        if not pages:
            logger.error(f"❌ No pages scraped after all attempts for KB {knowledge_base_id}")
            # Update KB status to error
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
            if kb:
                kb.status = "error"
                db.commit()
            return
        
        logger.info(f"✅ Successfully got {len(pages)} pages for processing")
        
        # Convert enhanced scraper format to vector store format
        processed_pages = []
        for page_dict in pages:
            processed_page = {
                "url": page_dict.get("final_url", page_dict.get("url", "")),
                "title": page_dict.get("title", ""),
                "content": page_dict.get("text", page_dict.get("content", "")),
                "scraped_at": datetime.now(timezone.utc).isoformat(),
                "framework": page_dict.get("framework", "unknown"),
                "word_count": page_dict.get("word_count", 0),
                "content_hash": page_dict.get("hash", ""),
                "status": page_dict.get("status", 200),
                **page_dict.get("meta", {})
            }
            
            # Only add pages with meaningful content
            content_length = len(processed_page["content"].strip()) if processed_page["content"] else 0
            if content_length > 50:
                processed_pages.append(processed_page)
                logger.info(f"📄 Processed: {processed_page['url']} ({content_length} chars, {processed_page.get('word_count', 0)} words)")
            else:
                logger.warning(f"⚠️ Skipping page with insufficient content: {processed_page['url']} (content length: {content_length})")
        
        if not processed_pages:
            logger.error(f"❌ No valid pages with content to process for KB {knowledge_base_id}")
            # Update KB status to error
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
            if kb:
                kb.status = "error"
                db.commit()
            return
        
        # Process pages through vector store
        logger.info(f"📄 Processing {len(processed_pages)} pages through vector store...")
        try:
            await vector_store.process_pages(processed_pages)
            total_chunks = vector_store.get_total_chunks()
            logger.info(f"✅ Vector store processing complete: {total_chunks} chunks created")
        except Exception as vector_error:
            logger.error(f"❌ Vector store processing failed: {vector_error}")
            # Update KB status to error
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
            if kb:
                kb.status = "error"
                db.commit()
            return
        
        # IMPORTANT: Validate final chunk count against subscription limits
        if total_chunks == 0:
            logger.error(f"❌ No chunks were created from {len(processed_pages)} pages")
            # Update KB status to error
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
            if kb:
                kb.status = "error"
                db.commit()
            return
        
        if not subscription.can_add_chunks(total_chunks):
            logger.error(f"❌ Chunk limit exceeded during processing: {total_chunks} chunks, user has {subscription.get_remaining_chunks()} remaining")
            
            # Clear the vector store to prevent limit violation
            try:
                vector_store.clear_data()
                logger.info(f"🧹 Cleared vector store due to chunk limit violation")
            except:
                pass
            
            # Update KB status to error
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
            if kb:
                kb.status = "error"
                db.commit()
            
            return
        
        # Update chunk usage in subscription
        subscription_manager.update_chunk_usage(user_id, total_chunks, db)
        logger.info(f"📊 Updated subscription chunk usage: +{total_chunks} chunks")
        
        # Verify vector store is ready
        if not vector_store.is_ready():
            logger.error(f"❌ Vector store is not ready after processing!")
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
            if kb:
                kb.status = "error"
                db.commit()
            return
        
        # Update knowledge base status
        kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
        if kb:
            kb.status = "ready"
            kb.total_chunks = total_chunks
            kb.last_updated = datetime.now(timezone.utc)
            db.commit()
            logger.info(f"✅ Updated KB status to ready with {total_chunks} chunks")
        
        logger.info(f"🎯 Website processing completed successfully: {len(processed_pages)} pages, {total_chunks} chunks")
        
    except Exception as e:
        logger.error(f"❌ Critical error in website processing: {str(e)}")
        import traceback
        logger.error(f"❌ Traceback: {traceback.format_exc()}")
        
        # Update KB status to error
        try:
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
            if kb:
                kb.status = "error"
                db.commit()
        except:
            pass
    finally:
        db.close()

# Usage validation endpoints
@app.get("/subscription/validate/chunks/{chunk_count}")
async def validate_chunk_usage(
    chunk_count: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Validate if user can add specified chunks"""
    validation = validate_chunk_addition(current_user.id, chunk_count, db)
    return ChunkUsageValidation(**validation)

@app.get("/subscription/validate/knowledge-base")
async def validate_kb_creation(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Validate if user can create a knowledge base"""
    validation = validate_kb_creation(current_user.id, db)
    return KnowledgeBaseValidation(**validation)

# Admin endpoint to recalculate usage
@app.post("/admin/recalculate-usage/{user_id}")
async def admin_recalculate_usage(
    user_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Admin endpoint to recalculate user's actual usage"""
    # Add admin role checking here if needed
    
    subscription_manager.recalculate_usage(user_id, db)
    
    return MessageResponse(message=f"Usage recalculated for user {user_id}")

# Backward compatibility: Keep old function name as alias
async def process_website_background(user_id: str, knowledge_base_id: str, config: WebsiteConfig):
    """Backward compatibility wrapper"""
    await process_website_background_enhanced(user_id, knowledge_base_id, config)

# Chat endpoints (unchanged)
@app.post("/chat/query", response_model=QueryResponse)
async def query_knowledge_base(
    query: QueryRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Verify ownership
    knowledge_base = db.query(KnowledgeBase).filter(
        KnowledgeBase.id == query.knowledge_base_id,
        KnowledgeBase.user_id == current_user.id
    ).first()
    
    if not knowledge_base:
        raise HTTPException(status_code=404, detail="Knowledge base not found")
    
    if knowledge_base.status != "ready":
        raise HTTPException(status_code=400, detail="Knowledge base is not ready")
    
    # Get vector store
    vector_store = vector_manager.get_vector_store(str(current_user.id), query.knowledge_base_id)
    
    if not vector_store.is_ready():
        raise HTTPException(status_code=400, detail="Vector store not ready")
    
    # Process query
    try:
        result = await vector_store.process_query(query.question, query.max_results)
        
        logger.info(f"Query processed for user {current_user.email}, KB {knowledge_base.name}")
        
        return QueryResponse(
            answer=result["answer"],
            sources=result.get("sources", []),
            confidence=result.get("confidence", 0.0)
        )
    except Exception as e:
        logger.error(f"Error processing query: {e}")
        raise HTTPException(status_code=500, detail="Error processing query")

@app.get("/knowledge-bases/{knowledge_base_id}/status")
async def get_knowledge_base_status(
    knowledge_base_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    knowledge_base = db.query(KnowledgeBase).filter(
        KnowledgeBase.id == knowledge_base_id,
        KnowledgeBase.user_id == current_user.id
    ).first()
    
    if not knowledge_base:
        raise HTTPException(status_code=404, detail="Knowledge base not found")
    
    return {
        "status": knowledge_base.status,
        "total_chunks": knowledge_base.total_chunks,
        "last_updated": knowledge_base.last_updated,
        "scraper_version": "enhanced_v4.1"
    }

# Widget helper function
def widget_to_response(widget: ChatWidget) -> dict:
    """Convert a ChatWidget database object to response format"""
    return {
        "id": widget.id,
        "name": widget.name,
        "knowledge_base_id": widget.knowledge_base_id,
        "widget_key": widget.widget_key,
        "website_url": widget.website_url,
        "primary_color": widget.primary_color,
        "widget_position": widget.widget_position,
        "welcome_message": widget.welcome_message,
        "placeholder_text": widget.placeholder_text,
        "widget_title": widget.widget_title,
        "is_active": widget.is_active,
        "show_branding": widget.show_branding,
        "allowed_domains": widget.allowed_domains,
        "total_conversations": widget.total_conversations,
        "total_messages": widget.total_messages,
        "last_used": widget.last_used.isoformat() if widget.last_used else None,
        "created_at": widget.created_at.isoformat() if widget.created_at else None,
    }

# Widget management endpoints (unchanged)
@app.post("/widgets", response_model=ChatWidgetResponse)
async def create_widget(
    widget_data: ChatWidgetCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Verify ownership of knowledge base
    knowledge_base = db.query(KnowledgeBase).filter(
        KnowledgeBase.id == widget_data.knowledge_base_id,
        KnowledgeBase.user_id == current_user.id
    ).first()
    
    if not knowledge_base:
        raise HTTPException(status_code=404, detail="Knowledge base not found")
    
    # Create widget
    widget = ChatWidget(
        user_id=current_user.id,
        knowledge_base_id=widget_data.knowledge_base_id,
        name=widget_data.name,
        website_url=widget_data.website_url,
        primary_color=widget_data.primary_color,
        widget_position=widget_data.widget_position,
        welcome_message=widget_data.welcome_message,
        placeholder_text=widget_data.placeholder_text,
        widget_title=widget_data.widget_title,
        show_branding=widget_data.show_branding,
        allowed_domains=','.join(widget_data.allowed_domains) if widget_data.allowed_domains else None
    )
    
    db.add(widget)
    db.commit()
    db.refresh(widget)
    
    logger.info(f"Widget created: {widget_data.name} for user {current_user.email}")
    
    return ChatWidgetResponse(**widget_to_response(widget))

@app.get("/widgets", response_model=List[ChatWidgetResponse])
async def list_widgets(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    widgets = db.query(ChatWidget).filter(
        ChatWidget.user_id == current_user.id
    ).all()
    
    return [ChatWidgetResponse(**widget_to_response(widget)) for widget in widgets]

@app.get("/widgets/{widget_id}", response_model=ChatWidgetResponse)
async def get_widget(
    widget_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    widget = db.query(ChatWidget).filter(
        ChatWidget.id == widget_id,
        ChatWidget.user_id == current_user.id
    ).first()
    
    if not widget:
        raise HTTPException(status_code=404, detail="Widget not found")
    
    return ChatWidgetResponse(**widget_to_response(widget))

@app.put("/widgets/{widget_id}", response_model=ChatWidgetResponse)
async def update_widget(
    widget_id: str,
    widget_data: ChatWidgetUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    widget = db.query(ChatWidget).filter(
        ChatWidget.id == widget_id,
        ChatWidget.user_id == current_user.id
    ).first()
    
    if not widget:
        raise HTTPException(status_code=404, detail="Widget not found")
    
    # Update fields
    for field, value in widget_data.dict(exclude_unset=True).items():
        if field == 'allowed_domains' and value is not None:
            setattr(widget, field, ','.join(value))
        else:
            setattr(widget, field, value)

    widget.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(widget)
    
    return ChatWidgetResponse(**widget_to_response(widget))

@app.delete("/widgets/{widget_id}")
async def delete_widget(
    widget_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    widget = db.query(ChatWidget).filter(
        ChatWidget.id == widget_id,
        ChatWidget.user_id == current_user.id
    ).first()
    
    if not widget:
        raise HTTPException(status_code=404, detail="Widget not found")
    
    # Delete associated conversations
    db.query(WidgetConversation).filter(WidgetConversation.widget_id == widget_id).delete()
    
    db.delete(widget)
    db.commit()
    
    return MessageResponse(message="Widget deleted successfully")

# Public widget endpoints (unchanged)
@app.get("/widget/{widget_key}/config", response_model=WidgetConfigResponse)
async def get_widget_config(widget_key: str, db: Session = Depends(get_db)):
    widget = db.query(ChatWidget).filter(
        ChatWidget.widget_key == widget_key,
        ChatWidget.is_active == True
    ).first()
    
    if not widget:
        raise HTTPException(status_code=404, detail="Widget not found or inactive")
    
    return WidgetConfigResponse(
        widget_key=widget.widget_key,
        primary_color=widget.primary_color,
        widget_position=widget.widget_position,
        welcome_message=widget.welcome_message,
        placeholder_text=widget.placeholder_text,
        widget_title=widget.widget_title,
        show_branding=widget.show_branding,
        is_active=widget.is_active
    )

@app.post("/widget/{widget_key}/chat", response_model=WidgetChatResponse)
async def widget_chat(
    widget_key: str,
    chat_request: WidgetChatRequest,
    db: Session = Depends(get_db)
):
    start_time = time_module.time()
    
    # Get widget
    widget = db.query(ChatWidget).filter(
        ChatWidget.widget_key == widget_key,
        ChatWidget.is_active == True
    ).first()
    
    if not widget:
        raise HTTPException(status_code=404, detail="Widget not found or inactive")
    
    # Generate session ID if not provided
    session_id = chat_request.session_id or str(uuid.uuid4())
    
    # Get vector store and process query
    try:
        vector_store = vector_manager.get_vector_store(widget.user_id, widget.knowledge_base_id)
        
        if not vector_store.is_ready():
            raise HTTPException(status_code=400, detail="Knowledge base not ready")
        
        result = await vector_store.process_query(chat_request.message, 5)
        
        # Log conversation
        conversation = WidgetConversation(
            widget_id=widget.id,
            session_id=session_id,
            user_message=chat_request.message,
            bot_response=result["answer"],
            confidence_score=result.get("confidence", 0.0),
            response_time_ms=int((time_module.time() - start_time) * 1000),
            user_ip=None,
            user_agent='',
            referrer_url=''
        )
        
        db.add(conversation)
        
        # Update widget analytics
        widget.total_messages += 1
        widget.last_used = datetime.now(timezone.utc)

        db.commit()
        
        return WidgetChatResponse(
            response=result["answer"],
            session_id=session_id,
            confidence=result.get("confidence", 0.0),
            sources=result.get("sources", [])
        )
        
    except Exception as e:
        logger.error(f"Error in widget chat: {e}")
        raise HTTPException(status_code=500, detail="Error processing message")

# Widget JavaScript endpoint (unchanged)
@app.get("/widget/{widget_key}/script.js")
async def get_widget_script(widget_key: str, db: Session = Depends(get_db)):
    widget = db.query(ChatWidget).filter(
        ChatWidget.widget_key == widget_key,
        ChatWidget.is_active == True
    ).first()
    
    if not widget:
        raise HTTPException(status_code=404, detail="Widget not found")
    
    # Generate the JavaScript widget code
    script = f"""
// AI Chatbot Widget - Generated for {widget.name} (Enhanced Scraper v4.1)
(function() {{
    var WIDGET_CONFIG = {{
        widgetKey: '{widget.widget_key}',
        apiUrl: '{os.getenv("WIDGET_API_URL", "http://localhost:8000")}',
        primaryColor: '{widget.primary_color}',
        position: '{widget.widget_position}',
        welcomeMessage: '{widget.welcome_message}',
        placeholderText: '{widget.placeholder_text}',
        title: '{widget.widget_title}',
        showBranding: {str(widget.show_branding).lower()},
        version: '4.1-enhanced'
    }};
    
    // Load widget CSS and JS
    var link = document.createElement('link');
    link.rel = 'stylesheet';
    link.href = WIDGET_CONFIG.apiUrl + '/static/widget.css';
    document.head.appendChild(link);
    
    var script = document.createElement('script');
    script.src = WIDGET_CONFIG.apiUrl + '/static/widget.js';
    script.onload = function() {{
        if (window.initChatWidget) {{
            window.initChatWidget(WIDGET_CONFIG);
        }}
    }};
    document.head.appendChild(script);
}})();
"""
    
    return Response(content=script, media_type="application/javascript")

# Document upload endpoints (unchanged for brevity - same as original)
@app.post("/knowledge-bases/{knowledge_base_id}/upload-documents", response_model=DocumentUploadResponse)
async def upload_documents(
    knowledge_base_id: str,
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload and process documents (TXT, DOCX) for a knowledge base"""
    
    # Verify ownership
    knowledge_base = db.query(KnowledgeBase).filter(
        KnowledgeBase.id == knowledge_base_id,
        KnowledgeBase.user_id == current_user.id
    ).first()
    
    if not knowledge_base:
        raise HTTPException(status_code=404, detail="Knowledge base not found")
    
    # Validate file types
    allowed_extensions = {'.txt', '.docx', '.doc'}
    file_data = []
    
    for file in files:
        if not file.filename:
            continue
            
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in allowed_extensions:
            continue
            
        # Check file size (max 10MB per file)
        content = await file.read()
        if len(content) > 10 * 1024 * 1024:  # 10MB limit
            continue
        
        # Store file data for background processing
        file_data.append({
            'filename': file.filename,
            'content': content,
            'file_ext': file_ext
        })
    
    if not file_data:
        raise HTTPException(
            status_code=400, 
            detail="No valid files found. Please upload TXT or DOCX files (max 10MB each)"
        )
    
    # Update status to processing
    knowledge_base.status = "processing"
    db.commit()
    
    # Process files in background
    background_tasks.add_task(
        process_documents_background,
        str(current_user.id),
        knowledge_base_id,
        file_data
    )
    
    logger.info(f"Document processing started for KB {knowledge_base_id}: {len(file_data)} files")
    
    return DocumentUploadResponse(
        message=f"Processing {len(file_data)} documents",
        knowledge_base_id=knowledge_base_id,
        status="processing",
        files_processed=0,
        total_chunks_added=0
    )

async def process_documents_background(user_id: str, knowledge_base_id: str, file_data: List[Dict]):
    """Background task to process uploaded documents"""
    from models import SessionLocal
    db = SessionLocal()
    
    # Use Docker temp directory
    temp_dir = os.environ.get('TMPDIR', '/app/tmp')
    
    try:
        # Get vector store for this user and knowledge base
        vector_store = vector_manager.get_vector_store(user_id, knowledge_base_id)
        processed_files = 0
        total_chunks_added = 0
        
        for file_info in file_data:
            try:
                # Create temp file in proper directory
                temp_file_path = os.path.join(temp_dir, f"upload_{os.getpid()}_{file_info['filename']}")
                
                # Write file content
                with open(temp_file_path, 'wb') as temp_file:
                    temp_file.write(file_info['content'])
                
                # Extract text based on file type
                text_content = extract_text_from_file(temp_file_path, file_info['filename'])
                
                if text_content.strip():
                    # Create document data for processing
                    document_data = {
                        "text": text_content,
                        "metadata": {
                            "source_url": f"uploaded_file://{file_info['filename']}",
                            "title": file_info['filename'],
                            "file_type": file_info['file_ext'],
                            "uploaded_at": datetime.now(timezone.utc).isoformat()
                        }
                    }
                    
                    # Process through vector store
                    chunks_added = await vector_store.process_document(document_data)
                    total_chunks_added += chunks_added
                    processed_files += 1
                    
                    logger.info(f"Processed document {file_info['filename']}: {chunks_added} chunks added")
                
                # Clean up temp file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                
            except Exception as e:
                logger.error(f"Error processing file {file_info['filename']}: {e}")
                # Clean up temp file if it exists
                if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
                    try:
                        os.unlink(temp_file_path)
                    except:
                        pass
                continue
        
        # Update knowledge base status
        kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
        if kb:
            kb.status = "ready" if processed_files > 0 else "error"
            kb.total_chunks = vector_store.get_total_chunks()
            kb.last_updated = datetime.now(timezone.utc)
            db.commit()
        
        logger.info(f"Document processing completed for KB {knowledge_base_id}: {processed_files} files, {total_chunks_added} chunks")
        
    except Exception as e:
        logger.error(f"Error processing documents for KB {knowledge_base_id}: {e}")
        # Update status to error
        kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).first()
        if kb:
            kb.status = "error"
            db.commit()
    finally:
        db.close()

def extract_text_from_file(file_path: str, filename: str) -> str:
    """Extract text content from various file types"""
    file_ext = Path(filename).suffix.lower()
    
    try:
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        elif file_ext in ['.docx', '.doc']:
            try:
                doc = docx.Document(file_path)
                text_content = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text.strip())
                
                return '\n\n'.join(text_content)
                
            except Exception as e:
                logger.error(f"Error processing Word document {filename}: {e}")
                return ""
        
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return ""
            
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        return ""

@app.get("/knowledge-bases/{knowledge_base_id}/processing-status")
async def get_processing_status(
    knowledge_base_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get the current processing status of a knowledge base"""
    knowledge_base = db.query(KnowledgeBase).filter(
        KnowledgeBase.id == knowledge_base_id,
        KnowledgeBase.user_id == current_user.id
    ).first()
    
    if not knowledge_base:
        raise HTTPException(status_code=404, detail="Knowledge base not found")
    
    return {
        "status": knowledge_base.status,
        "files_processed": 0,  # Could track this in database if needed
        "total_chunks_added": knowledge_base.total_chunks,
        "error_message": None if knowledge_base.status != "error" else "Processing failed",
        "scraper_version": "enhanced_v4.1"
    }

# Serve static widget files
@app.get("/static/{file_path:path}")
async def serve_static(file_path: str):
    """Serve static widget files"""
    from fastapi.responses import FileResponse
    import os
    
    static_dir = os.path.join(os.path.dirname(__file__), "static")
    file_location = os.path.join(static_dir, file_path)
    
    if os.path.exists(file_location):
        return FileResponse(file_location)
    else:
        raise HTTPException(status_code=404, detail="File not found")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)